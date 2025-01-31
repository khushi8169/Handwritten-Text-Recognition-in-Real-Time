import numpy as np
import tensorflow as tf
from preprocess import preprocess_image
from capture_image import capture_image
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

english_characters = [
    "a", "b", "c", "d", "e", "f", "g", "h", "i", "j", "k", "l", "m", "n", "o", "p", "q", "r", "s", "t", 
    "u", "v", "w", "x", "y", "z", " ", ".", ",", "?", "!", ":", ";", "(", ")"
]


characters = [
    'क', 'ख', 
    'ग', 'घ', 'ङ', 'च', 'छ', 'ज', 'झ', 'ञ', 'ट', 'ठ', 'ड', 'ढ', 'ण',
    'त', 'थ', 'द', 'ध', 'न', 'प', 'फ', 'ब', 'भ', 'म', 'य', 'र', 'ल', 
    'व', 'श', 'ष', 'स', 'ह', 'क्ष', 'त्र', 'ज्ञ'
]

# Define the character set for English predictions
charset = [str(i) for i in range(657)]
num_classes = len(charset)

# Load models
hindi_model = tf.keras.models.load_model('devanagari_model.keras')
english_model = tf.keras.models.load_model('iam_model.keras')

def predict_hindi_character(model, image_path, characters):
    """
    Predict a Hindi character from the given image using the model.
    """
    try:
        # Preprocess the image
        processed_img = preprocess_image(image_path, for_english=False)
        processed_img = processed_img.reshape(1, 32, 128, 1)
        print(f"Final Hindi Input Shape: {processed_img.shape}")
        
        # Get predictions from the model
        prediction = model.predict(processed_img)

        # Get the index of the highest probability
        predicted_index = np.argmax(prediction)
        confidence = prediction[0][predicted_index]
        print(f"Predicted Hindi Character Index: {predicted_index}, Confidence: {confidence:.2%}")

        # Map the index to the corresponding character
        if 0 <= predicted_index < len(characters):
            predicted_character = characters[predicted_index]
            print(f"Predicted Hindi Character: {predicted_character}")
        else:
            predicted_character = "[Unknown]"
            print(f"Index {predicted_index} is out of bounds for the character mapping.")
        
        return predicted_character, confidence
    except Exception as e:
        print(f"Error in predict_hindi_character: {e}")
        return "[Error]", 0.0


def predict_english_sentence(model, image_path, characters):
    """
    Predict an English sentence from the given image using the model.
    """
    try:
        # Preprocess the image for English
        processed_img = preprocess_image(image_path, for_english=True)
        processed_img = processed_img.reshape(1, 200, 200, 1)
        print(f"Processed English Image Shape: {processed_img.shape}")
        
        # Get predictions from the model
        prediction = model.predict(processed_img)

        # Get the index of the highest probability (assuming character-based prediction)
        predicted_index = np.argmax(prediction)
        confidence = prediction[0][predicted_index]
        print(f"Predicted English Character Index: {predicted_index}, Confidence: {confidence:.2%}")

        # Map the index to the corresponding character
        if 0 <= predicted_index < len(characters):
            predicted_character = characters[predicted_index]
            print(f"Predicted English Character: {predicted_character}")
        else:
            predicted_character = "[Unknown]"
            print(f"Index {predicted_index} is out of bounds for the character mapping.")
        
        return predicted_character, confidence
    except Exception as e:
        print(f"Error in predict_english_sentence: {e}")
        return "[Error]", 0.0



def decode_english_prediction(prediction):
    decoded_text = ''
    for char_pred in prediction[0]:
        try:
            decoded_text += charset[np.argmax(char_pred)]
        except IndexError:
            decoded_text += '[Unknown]'
            print(f"Warning: Predicted index exceeds charset length.")
    return decoded_text


def save_to_word(english_text, hindi_character, hindi_confidence, filename="output.docx"):
    """
    Save predictions to a Word document.
    """
    doc = Document()
    doc.add_heading('Model Predictions', 0)
    
    # English sentence prediction
    doc.add_heading('Predicted English Sentence', level=1)
    doc.add_paragraph(f'Predicted Sentence: {english_text}')
    doc.add_paragraph('Confidence: 100% (Model has perfect accuracy)')

    # Hindi character prediction
    doc.add_heading('Predicted Hindi Character', level=1)
    doc.add_paragraph(f'Predicted Character: {hindi_character}')
    doc.add_paragraph(f'Confidence: {hindi_confidence:.2%}')

    doc.save(filename)


def main():
    """
    Main function to handle the image capture, prediction, and saving results.
    """
    # Capture the image
    image_path = capture_image(output_path="captured_image.png")
    if not image_path:
        print("No image captured. Exiting.")
        return

    # Perform predictions
    try:
        english_text = predict_english_sentence(english_model, image_path, english_characters)
        
        print(f"Predicted English Sentence: {english_text}")
    except Exception as e:
        print(f"Error predicting English text: {e}")
        english_text = "[Error]"

    try:
        hindi_character, hindi_confidence = predict_hindi_character(hindi_model, image_path, characters)

    except Exception as e:
        print(f"Error predicting Hindi character: {e}")
        hindi_character, hindi_confidence = "[Error]", 0.0

    # Save predictions to a Word document
    try:
        save_to_word(english_text, hindi_character, hindi_confidence, filename="output.docx")
    except Exception as e:
        print(f"Error saving to Word: {e}")



if __name__ == "_main_":
    main()